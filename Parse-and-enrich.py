#!/usr/bin/python3

__author__ = 'jos_ir_'
__version__ = '1.5'
__date__ = 'Aug 3, 2022'

import ipaddress
import os
import argparse
import logging as log
import re
import time
import datetime
import operator
import socket
import json
import csv
from ipaddress import ip_address
from collections import defaultdict

try:
    from openpyxl import load_workbook
    import docx
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    import ipinfo
    import xlrd
    from pdfminer.high_level import extract_text
except Exception as e:
    log.error('Error loading libary: %s. Install libraries via pip3 install -r requirements.txt' %e)
    exit()

parser = argparse.ArgumentParser(description = 'This script looks for indicators in files with the extension xlsx, xls, docx, txt and enriches the IP with data from ipinfo (organization, geolocation etc). It outputs the data to a csv file.')
parser.add_argument('-i', '--inputpath', help='Select input file(s) that you want to search for IP addresses (accepted are: xlsx, xls, docx, txt). Exampe: -i path/* . *.csv is also accepted.', nargs='+', required=True)
parser.add_argument('-v', '--verbose', action='store_true')
parser.add_argument('-d', '--delay', help='ratelimit querying the API for x seconds', default=0, required=False)
parser.add_argument('-o', '--output', help='Output file.', required=False, default=f'{datetime.datetime.now():%Y-%m-%d_%H%M%S}_results.csv')
parser.add_argument('-a', '--accesstoken_location', help='Access token for ip_info (specify file. For example: --accesstoken_location custom_ip_info.key)', required=False, default='ip_info.key')
parser.add_argument('-s', '--search', help='Search for which indicators? Options: mobile phone number, e-mail address, url\'s, and IP addresses. Default is all. Example: --search ipaddress mobile', default=['ipaddress','mobile','email', 'url', 'md5', 'sha1', 'sha256', 'custom'] )
parser.add_argument('-se', '--skip_enrich', help='', action='store_false')
parser.add_argument('-csv_e', '--enrich_existing',  help='Add new columns with enriched ip_info data to existing CSV\'s (copies the original and adds the new columns there). Only works with csv files.', action='store_true')
parser.add_argument('-csv_q', '--quotechar', help='Quotecharacter vor reading and writing to csv files.', required=False,default='\"')
parser.add_argument('-csv_d', '--delimiter', help='Delimiter character for reading and writing to csv files.', required=False,default=',')
parser.add_argument('-csv_c', '--encoding', help='Encoding types for reading and writing to csv files. See encoding types: https://docs.python.org/3/library/codecs.html#standard-encodings', required=False,default='UTF-8')
args = parser.parse_args()

quotechar = args.quotechar
delimiter = args.delimiter
skip_enrich = args.skip_enrich
encoding = args.encoding
inputpath = args.inputpath
delay = args.delay
accesstoken_location = args.accesstoken_location
output = args.output
enrich_existing = args.enrich_existing
search = args.search
regex_result = {}
db_regex_result = {}
regex = {}
splitchars = '[,; \t]' #characters that are used to split values on the same line. This is needed to match (for example) 2 ip addresses in one line. For example: "8.8.8.8  9.9.9.9", becomes "8.8.8.8" and "9.9.9.9"
beginend = r'[ \t<>"\':;,.()]?' #characters that that the regexes to match for IP's, email addresses etc. For example: ">jos@ir.nl<" becomes "jos@ir.nl".
regex['url'] = r'(h(t|x)(t|x)ps?://(?:[-\w.]|(?:%[\da-fA-F]{2}))+)' #Regex from https://www.codegrepper.com/code-examples/python/match+url+regex+python
regex['email'] = r'([a-zA-Z0-9]+[\._]?[a-zA-Z0-9]+[@]\w+[.]\w{2,3})' #Regex from https://www.c-sharpcorner.com/article/how-to-validate-an-email-address-in-python/
regex['mobile'] = r'([\+]?[(]?[0-9]{2,3}[)]?[-\s\.]?[0-9]{3}[-\s\.]?[0-9]{4,6})' #Regex from https://ihateregex.io/expr/phone/
regex['ipaddress'] = r'([0-9]{1,3}\.[0-9]{1,3}\.[0-9]{1,3}\.[0-9]{1,3})' #Regex to match IP addresses. Note that the match will be further validated to check if it is truly a public ipv4 address.
regex['md5'] = r'([a-fA-F0-9]{32})'
regex['sha1'] = r'([a-fA-F0-9]{40})'
regex['sha256'] = r'([a-fA-F0-9]{64})'
regex['custom'] = r'(EnterYourCustomIndicatorHere)' #Regex to match a custom indicator

if args.verbose:
    log.basicConfig(format="%(levelname)s: %(message)s", level=log.INFO)
else:
    log.basicConfig(format="%(levelname)s: %(message)s")

def main(): 
    global db_regex_result
    db_regex_result = defaultdict(dict)

    if skip_enrich:
        
        try:
            accesstoken = open(accesstoken_location).readline().rstrip()
            log.info('Access token is: %s. and it was read from file: %s' %(accesstoken, accesstoken_location))

        except Exception as e:
            log.error('Could not load accesstoken: %s. Use -se (skip_enrich). Script will exit.' %e)
            exit()

        try:
            log.info('Setting up handler to enrich IP data with ipinfo.')
            global handler
            handler = ipinfo.getHandler(accesstoken)
            log.info('Testing the ability to query the api of ipinfo.io with 8.8.8.8')
            handler.getDetails('8.8.8.8.8')
            log.info('Test successful')
        except Exception as e:
            log.error('Query ipinfo.io api with 8.8.8.8 failed with error: %s. Use -se (skip_enrich). Script will exit.', e)
            exit()

    for file in inputpath:
        if not file.endswith('_enriched.csv'):
                if os.path.exists(file):
                    log.info('File %s exists! Whoooopwhooooop!' %file)
                else:
                    log.error('Path/file %s does not exist. With \'-i\' you can select the files that contain the indicators. For example: \'-i file_with_indicators.csv\' or \'-i *file.csv\'' %inputpath)
                    exit()
            
    for file in inputpath:
        if not file.endswith('_enriched.csv'):
            if file.endswith(('.txt', '.csv')):
                skip = False
                log.info('Reading txt or csv file: %s' %file)
                with open(file, 'r', encoding=encoding) as f:
                    reader = csv.reader(f)
                    try:
                        for line in reader:
                            break
                    except Exception as e:
                        log.error('Tried reading file %s but failed with error: %s. Script will continue without reading this file. Try converting the file to utf-8 or use (for example) --encoding UTF16' %(file,e))       
                        skip = True

                    if not skip:
                        for line in reader:
                            res = re.split(splitchars, str(line))
                            if res:
                                for word in res:
                                    regex_result = regex_search(word,search,file)
                                    if regex_result: 
                                        append_dictionary(regex_result,file,search)

            if file.endswith('.xlsx') and not file.startswith('~$'):
                log.info('Reading Excel (xlsx) document: %s' %file)
                wb = load_workbook(file)
                for sheet in wb.worksheets:
                    for row in sheet.iter_rows():
                        for cell in row:
                            regex_result = regex_search(str(cell.value),search,file)
                            if regex_result:
                                append_dictionary(regex_result,file,search)

            if file.endswith('.xls') and not file.startswith('~$'):
                log.info('Reading Excel (xls) document: %s' %file)
                book = xlrd.open_workbook(file)
                for i in range(book.nsheets):
                        try:
                            sh = book.sheet_by_index(i)
                        except:
                            pass
                        if sh:
                            for rx in range(sh.nrows):
                                data = sh.row(rx) 
                                if data:
                                    regex_result = regex_search(str(data),search,file)
                                    if regex_result:
                                        append_dictionary(regex_result,file,search)

            if file.endswith('.docx') and not file.startswith('~$'):
                log.info('Reading Word document: %s' %file)
                doc = docx.Document(file)
                paragraphs = [p.text for p in doc.paragraphs if p.text]
                for paragraph in paragraphs:
                    try:
                        res = re.split(splitchars, paragraph)
                    except:
                        pass
                    if res:
                        for word in res:
                            regex_result = regex_search(word,search,file)
                            if regex_result:
                                append_dictionary(regex_result,file,search)

                log.info('Reading tables from Word document: %s' %file)
                tables = doc.tables
                for index, table in enumerate(doc.tables):
                    for row in range(len(table.rows)):
                        for col in range(len(table.columns)):
                            try:
                                data = table.cell(row, col).text
                            except Exception:
                                pass

                            if data:
                                regex_result = regex_search(data,search,file)
                                if regex_result:
                                    append_dictionary(regex_result,file,search)

                log.info('Reading hyperlinks from Word document: %s' %file)
                rels = doc.part.rels
                for rel in rels:
                    if rels[rel].reltype == RT.HYPERLINK:
                        link = rels[rel]._target
                        regex_result = regex_search(link,search,file)
                        if regex_result:
                            append_dictionary(regex_result,file,search)

            if file.endswith('.pdf'):
                filehandle = open(file, 'r')
                log.info('Reading pdf file: %s' %file)
                text = extract_text(file)
                for word in repr(text).split(' '):
                    regex_result = regex_search(word.strip('\\n'), search,file)
                    if regex_result: 
                        append_dictionary(regex_result,file,search)

    if len(db_regex_result) == 0:
        log.error('Found nothing: Zero. Nada. Noppes. Nul. Nil. Niente. Ekkert. Faic. Res. Niks. Exiting script.')
        exit()

    log.info('Opening filehandle to file: %s' %output)
    with open(output, 'w', encoding=encoding, newline='') as f:
        header = ['Regex result', 'Count', 'Type', 'Found in file(s)', 'City', 'Country', 'Organization', 'Full', 'Error']
        writer = csv.writer(f)
        writer.writerow(header)
        print('{:<60} {:<10} {:<10}'.format('Regex result','Count', 'Type'))
        for key,value in db_regex_result.items():
            class result:
                regex = key
                count = value['count']
                type = value['type']
                found = value['Found in file(s)']
            if result.type == 'ipaddress':
                if skip_enrich:
                    ip_info = enrich(result.regex)
                    row = [result.regex, result.count, result.type, result.found, ip_info.city, ip_info.country, ip_info.org, json.dumps(ip_info.all), ip_info.error]
                    db_regex_result[result.regex] = ip_info.all
                else:
                    row = [result.regex, result.count, result.type, result.found]
                
            else:
                row = [result.regex, result.count, result.type, result.found, '', '', '', '', '']
            print('{:<60} {:<10} {:<10}'.format(result.regex, result.count, result.type))
            writer.writerow(row)

    if not skip_enrich:
        log.info("Script will not enrich with ipinfo.io data. Use the parameter: -csv_e (csv_enrich). And don't combine with -se (skip_enrich)")
        exit()

    if enrich_existing:
        log.info("Enriching existing csv files with ip_info data.")

        for file in inputpath:
            if file.endswith('.csv') and not file.endswith('_enriched.csv'):
                source = file
                target = source + '_enriched.csv'
                with open(source, 'r', encoding=encoding) as istr:
                    reader = csv.reader(istr, delimiter=delimiter, quotechar=quotechar)

                    with open(target, 'w', encoding=encoding, newline='') as ostr:
                        log.info('Reading csv file %s and copying it to %s and adding a column there with the enriched ip data.' %(source, target))

                        writer = csv.writer(ostr, delimiter=delimiter, quotechar=quotechar,quoting=csv.QUOTE_ALL)

                        for line in reader:
                            ip = regex_search(str(line),['ipaddress'],file)
                            if ip in db_regex_result:
                                line.append(json.dumps(db_regex_result[ip]))
                                writer.writerow(line)

                            else:
                                line.append('')
                                writer.writerow(line)


def append_dictionary(regex,file,search):
    if 'count' not in db_regex_result[regex]:
        db_regex_result[regex]['count'] = 1
    else:
        db_regex_result[regex]['count'] += 1
    if 'Found in file(s)' not in db_regex_result[regex]:
        db_regex_result[regex]['Found in file(s)'] = [file]
    elif file not in db_regex_result[regex]['Found in file(s)']:
        db_regex_result[regex]['Found in file(s)'].append(file)

def enrich(ip):
    log.info('Querying IP address against ipinfo: %s' %ip)
    remove_specialcharacters = r'[\'|"]'

    try:
        result = handler.getDetails(ip)
        setattr(result, 'error', '')

        if not hasattr(result,'org'):
            setattr(result, 'org', '')
        else:
            result.org = re.sub(remove_specialcharacters,r'', result.org)
        if not hasattr(result,'country'):
            setattr(result, 'country', '')
        else:
            result.country = re.sub(remove_specialcharacters,r'', result.country)
        if not hasattr(result,'city'):
            setattr(result, 'city', '')
        else:
            result.city = re.sub(remove_specialcharacters,r'', result.city)
        if not hasattr(result,'all'):
            setattr(result, 'all', '')
        else:
            for key,value in result.all.items():
                value = re.sub(remove_specialcharacters,r'', str(value))
                result.all[key] = value

    except Exception as e:
        log.error('Query %s against ipinfo.io resulted in an error: %s' %(ip,e))
        class result:
            org = ''
            country = ''
            city = '' 
            all = ''
            error = e
    time.sleep(int(delay))
    return result

def valid_ip(address):
    try: 
        if not ip_address(address).is_private:
            return address
    except:
        return False

def regex_search(value,search,full_path):
    for search_item in search:
        result = re.search(beginend + regex[search_item] + beginend, value)
        if result:
            result = result.group(1)
            if search_item == 'ipaddress':
                ip = result
                if valid_ip(ip):
                    log.info('Found the valid IP address \'%s\' in file \'%s\'.' %(ip,full_path))
                    db_regex_result[ip]['type'] = 'ipaddress'
                    return ip
                else:
                    log.info('IP address \'%s\' is not a valid public ipv4 address.' %result)

            if not search_item == 'ipaddress':
                log.info('Found the %s indicator \'%s\' in file \'%s\'.' %(search_item,result,full_path))        
                db_regex_result[result]['type'] = search_item
                return result

if __name__ == '__main__':
    main()